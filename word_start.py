import datetime
import random
import docx


# Поиск  данных
def search_object_in_src(the_object, src, add=False):
    print('src= ', src)
    doc = docx.Document(src)
    print('doc= ', doc)
    text = []
    the_paragraph = []
    for paragraph in doc.paragraphs:
        print('paragraph= ', paragraph)
        text.append(paragraph.text)
        print(f'{the_object} in? {paragraph.text}')
        if the_object.lower() in paragraph.text.lower():
            new_paragraph = paragraph.text.replace(the_object,
                                                   f"<b><i>{the_object}</i></b>")
            new_paragraph = new_paragraph.replace(the_object.lower(),
                                                  f"<b><i>{the_object.lower()}</i></b>")
            new_paragraph = new_paragraph.replace(the_object.upper(),
                                                  f"<b><i>{the_object.upper()}</i></b>")
            new_paragraph = new_paragraph.replace(the_object.capitalize(),
                                                  f"<b><i>{the_object.capitalize()}</i></b>")
            new_paragraph = new_paragraph.replace(the_object.title(),
                                                  f"<b><i>{the_object.title()}</i></b>")
            the_paragraph.append(new_paragraph)
    if not add:
        list_row = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if the_object.lower() in paragraph.text.lower():
                            list_row.append(row)
        if list_row:
            for row in list_row:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if the_object.lower() in paragraph.text.lower():
                            new_paragraph = paragraph.text.replace(the_object,
                                                                   f"<b><i>{the_object}</i></b>")
                            new_paragraph = new_paragraph.replace(the_object.lower(),
                                                                  f"<b><i>{the_object.lower()}</i></b>")
                            new_paragraph = new_paragraph.replace(the_object.upper(),
                                                                  f"<b><i>{the_object.upper()}</i></b>")
                            new_paragraph = new_paragraph.replace(the_object.capitalize(),
                                                                  f"<b><i>{the_object.capitalize()}</i></b>")
                            paragraph = new_paragraph.replace(the_object.title(),
                                                              f"<b><i>{the_object.title()}</i></b>")
                            the_paragraph.append(paragraph)
                        else:
                            the_paragraph.append(paragraph.text)

    print(text)
    print(the_paragraph)
    return the_paragraph


def get_date():
    td_ = datetime.date.today()
    td = td_.strftime('%d.%m.%Y')  # Разворачивает представление даты на европейский лад
    return td


# Заполняет шаблон
def fill_pattern(list_src, text):
    src_pattern = list_src[0]
    date = get_date()
    doc = docx.Document(src_pattern)
    dict_replace = {'*date': date, '*text': text}
    for key in dict_replace:

        for paragraph in doc.paragraphs:
            if key in paragraph.text:
                paragraph.text = dict_replace[key]

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, dict_replace[key])

    new_src = '/'.join(src_pattern.split('/')[:-1]) + '/' + src_pattern.split('/')[-1][16:]
    new_src = '.'.join(new_src.split('.')[:-1]) + '_' + date + '_' + str(random.randrange(10000)) + '.docx'
    doc.save(new_src)
    return new_src
